from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path.cwd() / "interview_owner_execution_guide.docx"

BLUE = RGBColor(31, 77, 121)
DARK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F8"
BORDER = "D9E2F3"


def set_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=DARK, size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_borders(table, color="DADCE0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_font(run, size=15 if level == 1 else 12.5, color=BLUE, bold=True)


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.18
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_font(lead, size=10.5, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_font(rest, size=10.5)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=10.5)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=10.5)


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=10.5)


def add_callout(doc, title, lines, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, BORDER)
    set_table_width(table, [Inches(6.35)])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_font(run, size=10.5, color=BLUE, bold=True)
    for line in lines:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.12
        run = paragraph.add_run(line)
        set_font(run, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_width(table, widths)

    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        shade_cell(cell, header_fill)
        set_cell_text(cell, header, bold=True, color=BLUE, size=9.8)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_checklist(doc, items):
    for item in items:
        add_bullet(doc, "□ " + item)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["List Bullet", "List Number"]:
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles[style_name].font.size = Pt(10.5)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("采访负责人执行手册")
    set_font(run, size=22, color=BLUE, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run("宿舍生活用品智能分类与补货提醒系统")
    set_font(run, size=12, color=GRAY)

    add_callout(
        doc,
        "这份文档只回答一件事：采访负责人到底要干什么。",
        [
            "你的工作不是“问几个问题”，而是把真实用户需求、真实样本、真实试用记录都收集齐。",
            "最终要让报告和视频能证明：用户真实存在、需求真实存在、系统真实交付过。",
        ],
        fill="F7FAFF",
    )

    add_heading(doc, "一、最终要交付什么")
    add_table(
        doc,
        ["材料", "要交的内容", "最低要求"],
        [
            ["用户确认", "目标用户身份、联系方式、是否愿意配合", "用户合规，愿意访谈、试用、留联系方式"],
            ["访谈记录", "背景、痛点、现有方法、系统期待", "有具体经历，不写空话"],
            ["真实物品样本", "用户实际生活用品数据", "至少 15 条，推荐 20-30 条"],
            ["测试记录", "用户输入、系统输出、用户评价", "至少 5 个正常案例，1 个异常或边界案例"],
            ["交付材料", "交付视频、用户反馈、回执单", "视频里能看出用户实际操作"],
            ["AI 使用记录", "AI 如何帮忙、哪里有问题、怎么修改", "说明真实用户内容不能由 AI 代替"],
        ],
        [Inches(1.25), Inches(3.15), Inches(1.75)],
    )

    add_heading(doc, "二、按顺序做这 6 步")
    for title, detail in [
        ("第 1 步：找用户", "找一个非本组、非本课程期末考核学生、非直系亲属的真实用户。优先找住校或独立管理生活用品的同学。"),
        ("第 2 步：确认配合", "提前确认对方愿意访谈、提供物品样本、后续试用系统、录制交付过程、留下联系方式、填写回执单。"),
        ("第 3 步：做首次访谈", "先问真实经历，再问现有解决方式，最后才问系统功能。不要一开始就引导用户说“我需要智能补货系统”。"),
        ("第 4 步：收集真实样本", "让用户列出自己真实使用的生活用品，记录物品名称、剩余量、使用频率、购买时间、用户判断等。"),
        ("第 5 步：同步给组员", "把用户样本和需求交给算法负责人，把用户希望看到的输入输出交给系统负责人。"),
        ("第 6 步：组织交付测试", "系统做好后让用户实际输入样本，记录系统结果、用户评价、不准确的地方和改进建议。"),
    ]:
        add_body(doc, title + "：" + detail, bold_lead=title)

    add_heading(doc, "三、找用户时先确认这些条件")
    add_table(
        doc,
        ["必须满足", "不要选择"],
        [
            ["不是本小组成员", "本组同学或帮忙演示的人"],
            ["不是参与本课程期末考核的学生", "同课程同学，容易被认为不合规"],
            ["不是小组成员直系亲属", "父母、兄弟姐妹等直系亲属"],
            ["有住校或独立管理生活用品经历", "完全没有相关生活场景的人"],
            ["愿意试用系统并留下联系方式", "不愿留联系方式或不愿录制交付过程的人"],
        ],
        [Inches(3.15), Inches(3.2)],
    )
    add_callout(
        doc,
        "联系用户可以这样说",
        [
            "你好，我们机器学习课程期末要做一个真实用户项目，方向是“宿舍生活用品智能分类与补货提醒系统”。主要想解决住校学生经常忘记补纸巾、洗衣液、牙膏、垃圾袋、药品、学习用品等问题。",
            "我们想请你作为目标用户，先做一次简单访谈，后面系统做好后再请你实际试用一次。过程中会记录你的需求和反馈，也需要你留下联系方式用于课程抽检。不会公开你的隐私信息。你愿意配合吗？",
        ],
    )

    add_heading(doc, "四、首次访谈怎么问")
    add_body(doc, "访谈顺序要自然：先问生活背景，再问真实痛点，再问现在怎么解决，最后才讨论系统需求。重点是让用户讲具体经历。")
    interview_sections = [
        (
            "1. 背景问题",
            [
                "你现在是住校吗？平时生活用品主要由谁购买和管理？",
                "你一般多久采购一次生活用品？更习惯网购还是去超市买？",
                "哪些东西是你自己用，哪些是和室友共用？",
                "你平时会固定记录生活用品剩余情况吗？",
                "课程、社团、考试或实习安排会不会影响你管理这些小事？",
            ],
        ),
        (
            "2. 痛点问题",
            [
                "你最近一次发现生活用品快用完，或者已经用完，是什么时候？",
                "当时缺的是什么东西？给你带来了什么不方便？",
                "哪些物品最容易被你忽略？为什么容易忘？",
                "有没有临时找室友借东西、临时下单、加急购买的经历？",
                "考试周、实习前、放假返校前，这种情况会不会更明显？",
            ],
        ),
        (
            "3. 现有解决方式问题",
            [
                "你现在一般怎么记这些东西？靠脑子、备忘录、购物车，还是室友提醒？",
                "你觉得现在这种方式最大的问题是什么？",
                "你有没有尝试过清单、备忘录或闹钟提醒？为什么没有坚持？",
                "你更希望主动记录，还是系统根据简单输入帮你判断？",
            ],
        ),
        (
            "4. 系统需求问题",
            [
                "如果有一个系统帮你管理宿舍生活用品，你最希望它帮你做什么？",
                "你愿意输入哪些信息？哪些信息太麻烦、不想填？",
                "你希望系统把物品分成哪些类别？",
                "你希望系统最后输出什么？例如分类、补货建议、紧急程度、预计还能用几天。",
                "你能接受“急需补货、近期关注、暂不需要、保质期提醒”这样的等级吗？",
                "如果系统判断错了，你希望能不能手动修改？",
            ],
        ),
        (
            "5. 后续交付确认问题",
            [
                "系统完成后，你是否愿意用自己的几件生活用品做一次测试？",
                "你是否愿意录制一段实际使用过程？可以不露脸，只录手部、屏幕和声音。",
                "你是否愿意填写交付回执单？",
                "你是否愿意留下联系方式，便于课程老师或助教后续抽检？",
            ],
        ),
    ]
    for title, questions in interview_sections:
        add_heading(doc, title, 2)
        for question in questions:
            add_number(doc, question)

    add_heading(doc, "五、访谈时一定要记录什么")
    add_table(
        doc,
        ["记录项", "怎么写才合格"],
        [
            ["用户背景", "写清楚是否住校、生活用品谁管理、为什么符合项目场景"],
            ["具体痛点", "写具体经历，例如“考试周发现纸巾没了，只能临时找室友借”"],
            ["现有方法", "写用户现在靠什么解决，以及为什么不好用"],
            ["功能期待", "写用户真正想要的输入、输出和提醒方式"],
            ["用户原话", "摘 2-4 句能体现痛点或需求的话"],
            ["采访人总结", "用自己的话说明这个用户为什么适合作为目标用户"],
        ],
        [Inches(1.55), Inches(4.8)],
    )

    add_heading(doc, "六、真实物品样本怎么收集")
    add_body(doc, "至少收集 15 条，推荐 20-30 条。不要全部让 AI 编。真实用户样本是报告里证明需求真实性的重要证据。")
    add_table(
        doc,
        ["字段", "示例", "说明"],
        [
            ["物品名称", "抽纸", "用户真实使用的物品"],
            ["用户描述", "宿舍每天用，只剩一小包", "尽量保留生活化描述"],
            ["用户类别", "清洁日用", "由用户或采访人初步标注"],
            ["购买日期/已使用天数", "2026-05-18 / 14 天", "二选一也可以"],
            ["剩余量", "20%", "如果用户说不清，可估算"],
            ["使用频率", "每天 / 每周 2 次", "用于补货预测"],
            ["使用人数", "1 人 / 2 人共用", "共用品要记录"],
            ["保质期", "无 / 2026-08-01", "药品、口罩、食品类要注意"],
            ["用户判断", "急需补货", "作为测试和标签参考"],
        ],
        [Inches(1.35), Inches(1.75), Inches(2.9)],
    )
    add_body(doc, "推荐类别：洗漱用品、清洁日用、学习用品、药品健康、电子配件、其他用品。")
    add_body(doc, "推荐补货等级：急需补货、近期关注、暂不需要、保质期提醒。")

    add_heading(doc, "七、交给算法和系统负责人的信息")
    add_table(
        doc,
        ["交给谁", "交什么"],
        [
            ["算法负责人", "真实物品样本、用户认为的类别、补货等级、典型测试案例、用户对“急需/关注/不需要”的理解"],
            ["系统负责人", "用户愿意填写的输入项、用户希望看到的输出项、必须支持的异常输入提示"],
        ],
        [Inches(1.45), Inches(4.9)],
    )
    add_callout(
        doc,
        "注意",
        [
            "系统不能只做成 if-else 规则。比如“剩余量低于 20% 就提醒”只能作为参考，不能成为唯一逻辑。",
            "你收集的数据要帮助小组体现两种机器学习任务：物品类别分类、补货紧急程度预测。",
        ],
    )

    add_heading(doc, "八、系统做好后怎么组织用户测试")
    add_body(doc, "测试时让用户亲自操作，不要只看小组成员演示。测试要包含正常案例，也要包含异常或边界案例。")
    add_table(
        doc,
        ["测试物品", "输入内容", "观察重点"],
        [
            ["抽纸", "买了两周，剩余 20%，每天使用", "是否识别为清洁日用，并提示较紧急"],
            ["洗衣液", "买了一个月，剩余 30%，每周洗衣 2 次", "是否提示近期关注或合理说明"],
            ["牙膏", "买了三周，剩余 40%，每天使用 2 次", "是否识别为洗漱用品"],
            ["感冒药", "剩余较多，保质期还有 2 个月", "是否识别为药品健康，并提醒保质期"],
            ["数据线", "使用时间较长，外皮破损，暂无备用", "是否识别为电子配件，并给出建议"],
            ["异常输入", "牙膏，剩余 130%", "是否提示输入异常"],
            ["信息不完整", "纸巾，只输入名称", "是否要求补充信息或给出有限判断"],
        ],
        [Inches(1.15), Inches(3.2), Inches(2.0)],
    )

    add_heading(doc, "九、测试记录必须写清楚")
    add_checklist(
        doc,
        [
            "用户输入了什么。",
            "系统输出了什么分类结果。",
            "系统输出了什么补货建议或紧急程度。",
            "用户觉得准不准，为什么。",
            "用户有没有修改输入。",
            "有没有出现异常输入或系统提示。",
            "用户提出了哪些改进建议。",
        ],
    )

    add_heading(doc, "十、交付视频怎么录")
    add_checklist(
        doc,
        [
            "开头说明这是给哪位真实用户交付，用户有什么需求。",
            "用户打开或进入系统。",
            "用户亲自输入至少 5 个真实生活用品案例。",
            "系统展示分类结果、补货建议、紧急程度。",
            "至少展示 1 个异常或边界输入。",
            "用户说出 1-2 条评价或改进建议。",
            "如果用户不方便露脸，可以录手部操作、屏幕和声音。",
        ],
    )

    add_heading(doc, "十一、交付回执单要填什么")
    add_table(
        doc,
        ["项目", "填写要求"],
        [
            ["交付时间", "写具体日期和时间"],
            ["交付对象", "写用户姓名或可识别称呼"],
            ["目标用户背景", "说明为什么这个用户符合项目场景"],
            ["交付步骤", "说明如何介绍系统、如何让用户输入、如何收集反馈"],
            ["用户联系方式", "必须真实，便于抽检"],
            ["用户评价", "既写满意点，也写不足和建议"],
            ["满意度", "可以用 1-5 分或文字描述"],
            ["用户签字", "纸质版或电子确认都要保存"],
        ],
        [Inches(1.45), Inches(4.9)],
    )

    add_heading(doc, "十二、最容易扣分的坑")
    for item in [
        "用户不合规：是本课程同学、小组成员或直系亲属。",
        "只有访谈，没有用户实际使用记录。",
        "视频只有小组成员演示，看不出真实用户上手。",
        "访谈像提前写好的剧本，没有具体生活经历。",
        "交付回执单没有联系方式或签字。",
        "数据全部由 AI 生成，没有真实用户样本。",
        "系统只是规则判断，没有体现机器学习模型。",
        "没有异常输入测试。",
        "用户反馈只有夸奖，没有不足和改进建议。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十三、完成前自查清单")
    add_checklist(
        doc,
        [
            "已找到合规目标用户。",
            "已完成首次需求访谈。",
            "已整理访谈记录和用户原话。",
            "已收集至少 15 条真实物品样本。",
            "已把用户需求同步给算法负责人和系统负责人。",
            "已准备正常、异常、边界测试用例。",
            "已组织用户实际试用系统。",
            "已录制交付视频素材。",
            "已整理用户反馈。",
            "已填写交付回执单。",
            "已保存 AI 辅助访谈设计和修改记录。",
        ],
    )
    add_callout(
        doc,
        "一句话记住",
        ["采访负责人负责的是“真实用户证据链”：找用户、问清楚、收样本、让用户试用、留下反馈和回执。"],
        fill="FFF8E8",
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("机器学习课程期末项目 | 采访负责人执行手册")
        set_font(run, size=9, color=GRAY)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
